from __future__ import annotations

import io
import os
import re
from pathlib import Path

from flask import jsonify, request

from services import codex_group_chat
from storage import r2_store


_MAX_IMPORT_BYTES = 16 * 1024 * 1024
_MAX_IMPORT_CHUNKS = 10
_IMPORT_CHUNK_TARGET_CHARS = 9000
_IMPORT_CHUNK_MAX_CHARS = 12000
_MAX_EXTRACT_CHARS = _MAX_IMPORT_CHUNKS * _IMPORT_CHUNK_MAX_CHARS
_TEXT_EXTS = {".txt", ".md", ".markdown"}
_PDF_EXTS = {".pdf"}
_WORD_EXTS = {".docx"}
_SOURCE_TYPE_OVERRIDES = {"pdf", "question_bank", "word", "text", "fenbi", "note", "wrong_question"}
_LOCAL_STUDY_ROOT = Path(os.environ.get("STUDYROOM_LOCAL_IMPORT_DIR") or (Path.home() / "Downloads" / "study")).expanduser()
_LOCAL_STUDY_MAX_CHARS = 800_000
_LOCAL_STUDY_MAX_PDF_PAGES = 500
_LOCAL_STUDY_MATERIAL_MAX_CHUNKS = 24
_LOCAL_STUDY_QUESTION_MAX_CHUNKS = 160
_LOCAL_STUDY_QUESTION_GROUP_SIZE = 20
_LOCAL_STUDY_EXTS = _TEXT_EXTS | _PDF_EXTS | _WORD_EXTS


def _clip_import_text(text: str) -> str:
    clean = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(clean) <= _MAX_EXTRACT_CHARS:
        return clean
    return clean[:_MAX_EXTRACT_CHARS].rstrip() + "\n\n...[导入时已截断，原文件仍可重新拆分导入]"


def _is_import_heading(line: str) -> bool:
    clean = str(line or "").strip()
    if not clean or len(clean) > 90:
        return False
    if re.match(r"^---\s*第\s*\d+\s*页\s*---$", clean):
        return False
    patterns = (
        r"^第[一二三四五六七八九十百千\d]+[章节讲课部分单元]\s*[:：、.\s-]?.{0,60}$",
        r"^[一二三四五六七八九十]+[、.．]\s*\S.{0,60}$",
        r"^\d+(?:\.\d+){0,3}[、.．\s]\s*\S.{0,60}$",
        r"^（[一二三四五六七八九十\d]+）\s*\S.{0,60}$",
    )
    return any(re.match(pattern, clean) for pattern in patterns)


def _chunk_label(label: str, fallback: str = "正文") -> str:
    clean = re.sub(r"\s+", " ", str(label or "").strip())
    clean = re.sub(r"^---\s*|\s*---$", "", clean).strip()
    return (clean or fallback)[:36]


def _split_sections_by_headings(text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_label = "开头"
    current_lines: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if _is_import_heading(line) and current_lines:
            body = "\n".join(current_lines).strip()
            if body:
                sections.append((current_label, body))
            current_label = line.strip()
            current_lines = [line]
        else:
            current_lines.append(line)
            if _is_import_heading(line):
                current_label = line.strip()
    body = "\n".join(current_lines).strip()
    if body:
        sections.append((current_label, body))
    return sections


def _split_sections_by_pages(text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    parts = re.split(r"(?=^---\s*第\s*\d+\s*页\s*---$)", text, flags=re.MULTILINE)
    for index, part in enumerate(parts, start=1):
        body = part.strip()
        if not body:
            continue
        first_line = body.splitlines()[0].strip()
        label = first_line if re.match(r"^---\s*第\s*\d+\s*页\s*---$", first_line) else f"第 {index} 段"
        sections.append((label, body))
    return sections or [("正文", text.strip())]


def _split_long_section(label: str, body: str) -> list[tuple[str, str]]:
    clean = str(body or "").strip()
    if len(clean) <= _IMPORT_CHUNK_MAX_CHARS:
        return [(label, clean)] if clean else []
    out: list[tuple[str, str]] = []
    current: list[str] = []
    for paragraph in re.split(r"\n\s*\n", clean):
        part = paragraph.strip()
        if not part:
            continue
        while len(part) > _IMPORT_CHUNK_MAX_CHARS:
            head = part[:_IMPORT_CHUNK_MAX_CHARS].rstrip()
            if current:
                out.append((label, "\n\n".join(current).strip()))
                current = []
            out.append((label, head))
            part = part[_IMPORT_CHUNK_MAX_CHARS:].lstrip()
        candidate = "\n\n".join([*current, part]).strip()
        if current and len(candidate) > _IMPORT_CHUNK_MAX_CHARS:
            out.append((label, "\n\n".join(current).strip()))
            current = [part]
        else:
            current.append(part)
    if current:
        out.append((label, "\n\n".join(current).strip()))
    return out


def _split_import_chunks(text: str, max_chunks: int = _MAX_IMPORT_CHUNKS) -> list[dict]:
    clean = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not clean:
        return []
    if len(clean) <= _IMPORT_CHUNK_MAX_CHARS:
        return [{"label": "正文", "content": clean}]

    sections = _split_sections_by_headings(clean)
    if len(sections) <= 1:
        sections = _split_sections_by_pages(clean)

    pieces: list[tuple[str, str]] = []
    for label, body in sections:
        pieces.extend(_split_long_section(label, body))

    chunks: list[dict] = []
    current_parts: list[str] = []
    current_labels: list[str] = []
    for label, body in pieces:
        candidate = "\n\n".join([*current_parts, body]).strip()
        if current_parts and len(candidate) > _IMPORT_CHUNK_TARGET_CHARS:
            chunks.append({"label": _chunk_label(current_labels[0] if current_labels else "正文"), "content": "\n\n".join(current_parts).strip()})
            current_parts = [body]
            current_labels = [label]
        else:
            current_parts.append(body)
            if label and label not in current_labels:
                current_labels.append(label)
    if current_parts:
        chunks.append({"label": _chunk_label(current_labels[0] if current_labels else "正文"), "content": "\n\n".join(current_parts).strip()})

    if max_chunks > 0 and len(chunks) > max_chunks:
        chunks = chunks[:max_chunks]
        chunks[-1]["content"] = (
            str(chunks[-1].get("content") or "").rstrip()
            + "\n\n...[后续内容超过自动拆分上限，先保留前面重点段落]"
        )
    return [chunk for chunk in chunks if str(chunk.get("content") or "").strip()]


def _chunked_title(title: str, index: int, total: int, label: str) -> str:
    clean_title = str(title or "").strip() or "未命名资料"
    if total <= 1:
        return clean_title
    clean_label = _chunk_label(label, f"第 {index} 段")
    return f"{clean_title}（{index}/{total}：{clean_label}）"


def _decode_text_file(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _has_enough_import_text(parts: list[str]) -> bool:
    return sum(len(part) for part in parts) >= _MAX_EXTRACT_CHARS


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("服务器缺少 pypdf，请先安装 requirements.txt") from exc

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for index, page in enumerate(reader.pages[:200], start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_text = page_text.strip()
        if page_text:
            parts.append(f"--- 第 {index} 页 ---\n{page_text}")
        if _has_enough_import_text(parts):
            break
    return "\n\n".join(parts)


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("服务器缺少 python-docx，请先安装 requirements.txt") from exc

    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text.strip())
        if _has_enough_import_text(parts):
            return "\n".join(parts)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
            if _has_enough_import_text(parts):
                return "\n".join(parts)
    return "\n".join(parts)


def _clip_local_import_text(text: str, max_chars: int = _LOCAL_STUDY_MAX_CHARS) -> str:
    clean = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(clean) <= max_chars:
        return clean
    return clean[:max_chars].rstrip() + "\n\n...[本地导入已截断，可单独加大上限重新导入]"


def _extract_local_pdf_text(path: Path, max_chars: int = _LOCAL_STUDY_MAX_CHARS, max_pages: int = _LOCAL_STUDY_MAX_PDF_PAGES) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("服务器缺少 pypdf，请先安装 requirements.txt") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    total = 0
    for index, page in enumerate(reader.pages[:max_pages], start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_text = page_text.strip()
        if page_text:
            block = f"--- 第 {index} 页 ---\n{page_text}"
            parts.append(block)
            total += len(block)
        if total >= max_chars:
            break
    return _clip_local_import_text("\n\n".join(parts), max_chars)


def _extract_local_docx_text(path: Path, max_chars: int = _LOCAL_STUDY_MAX_CHARS) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("服务器缺少 python-docx，请先安装 requirements.txt") from exc

    doc = Document(str(path))
    parts: list[str] = []
    total = 0
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
            total += len(text)
        if total >= max_chars:
            break
    return _clip_local_import_text("\n".join(parts), max_chars)


def _extract_local_file_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in _TEXT_EXTS:
        return _clip_local_import_text(_decode_text_file(path.read_bytes())), "text"
    if suffix in _PDF_EXTS:
        return _extract_local_pdf_text(path), "pdf"
    if suffix in _WORD_EXTS:
        return _extract_local_docx_text(path), "word"
    raise ValueError("暂只支持 pdf、docx、txt、md")


def _extract_import_text(filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename or "").suffix.lower()
    if suffix in _TEXT_EXTS:
        return _clip_import_text(_decode_text_file(content)), "text"
    if suffix in _PDF_EXTS:
        return _clip_import_text(_extract_pdf_text(content)), "pdf"
    if suffix in _WORD_EXTS:
        return _clip_import_text(_extract_docx_text(content)), "word"
    raise ValueError("暂只支持 pdf、docx、txt、md")


def _split_answer_section(text: str) -> tuple[str, str]:
    clean = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    matches = list(re.finditer(r"(?:^|\n)\s*(?:参考答案|答案速查|参考答案及解析|答案解析|答案：|答案)\s*", clean))
    if not matches:
        return clean, ""
    match = next((m for m in matches if m.start() >= len(clean) * 0.35), None)
    if not match:
        return clean, ""
    if match.start() < len(clean) * 0.35:
        return clean, ""
    return clean[: match.start()].strip(), clean[match.start() :].strip()


def _parse_answer_entries(answer_text: str) -> dict[str, str]:
    normalized = re.sub(r"【\s*答案\s*】", "答案", str(answer_text or ""))
    normalized = normalized.replace("正确答案", "答案")
    normalized = re.sub(r"正\s*确", "正确", normalized)
    normalized = re.sub(r"错\s*误", "错误", normalized)
    buckets: dict[str, set[str]] = {}
    answer_token = r"([A-H]{1,6}|正确|错误|对|错|√|✓|×|✕|X|T|F|TRUE|FALSE)"
    patterns = (
        rf"(?:^|[\s。；;，,])(?:第\s*)?(\d{{1,4}})\s*(?:题)?\s*[.、:：）)]?\s*(?:答案\s*[:：]?)?\s*{answer_token}(?=$|[\s。；;，,【\[])",
        rf"(?:^|\n)\s*(\d{{1,4}})\s*(?:题)?\s*答案\s*[:：]?\s*{answer_token}",
    )
    for pattern in patterns:
        for match in re.finditer(pattern, normalized, flags=re.IGNORECASE):
            number = str(match.group(1) or "").strip()
            raw_answer = str(match.group(2) or "").strip().upper()
            if raw_answer in {"正确", "对", "√", "✓", "T", "TRUE"}:
                answer = "A"
            elif raw_answer in {"错误", "错", "×", "✕", "X", "F", "FALSE"}:
                answer = "B"
            else:
                answer = re.sub(r"[^A-H]", "", raw_answer)
            if number and answer:
                buckets.setdefault(number, set()).add(answer)
    return {number: next(iter(values)) for number, values in buckets.items() if len(values) == 1}


def _question_numbers(text: str) -> list[str]:
    seen: list[str] = []
    for match in re.finditer(r"(?:^|\n)\s*(\d{1,4})\s*[、.．)]\s*(?=\S)", str(text or "")):
        number = str(match.group(1) or "").strip()
        if number and number not in seen:
            seen.append(number)
    return seen


def _is_question_bank_chapter_heading(line: str) -> bool:
    clean = _chunk_label(line, "")
    if not clean or len(clean) > 80:
        return False
    if re.search(r"\.{4,}|…{2,}", clean):
        return False
    if re.match(r"^---\s*第\s*\d+\s*页\s*---$", clean):
        return False
    if re.match(r"^\d{1,4}\s*[、.．)]\s*\S", clean):
        return False
    if re.match(r"^[A-H]\s*[.．、)]\s*", clean, flags=re.IGNORECASE):
        return False
    if re.match(r"^第[一二三四五六七八九十百千万零〇\d]+[章节编篇部分单元讲]\s*[:：、.\s-]?.{0,50}$", clean):
        return True
    if re.match(r"^(?:专题|模块|单元|章节)\s*[一二三四五六七八九十百千万零〇\d]+[：:、.\s-]?.{0,50}$", clean):
        return True
    if re.match(r"^(?:法理学|宪法|民法|刑法|行政法|经济法|商法|诉讼法|其他法律法规|党建|时政|三农|乡村振兴|基层治理|村务管理|公文写作|计算机)(?:\s|$|[：:、-]).{0,32}$", clean):
        return True
    return False


def _split_question_bank_chapters(text: str) -> list[dict]:
    clean = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not clean:
        return []
    chapters: list[dict] = []
    current_title = "未分章"
    current_lines: list[str] = []

    def flush() -> None:
        content = "\n".join(current_lines).strip()
        if content:
            chapters.append({"title": _chunk_label(current_title, "未分章"), "content": content})

    for raw in clean.splitlines():
        line = raw.strip()
        if _is_question_bank_chapter_heading(line):
            if _chapter_titles_match(current_title, line):
                current_lines.append(raw)
                continue
            flush()
            current_title = line
            current_lines = [line]
            continue
        current_lines.append(raw)
    flush()
    return chapters or [{"title": "未分章", "content": clean}]


def _split_question_blocks(text: str) -> tuple[str, list[dict]]:
    intro: list[str] = []
    blocks: list[dict] = []
    current_number = ""
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_number, current_lines
        content = "\n".join(current_lines).strip()
        if current_number and content:
            blocks.append({"number": current_number, "content": content})
        current_number = ""
        current_lines = []

    for raw in str(text or "").splitlines():
        match = re.match(r"^\s*(\d{1,4})\s*[、.．)]\s*(?=\S)", raw)
        if match:
            flush()
            current_number = str(match.group(1) or "").strip()
            current_lines = [raw]
        elif current_number:
            current_lines.append(raw)
        else:
            intro.append(raw)
    flush()
    return "\n".join(intro).strip(), blocks


def _normalize_chapter_key(text: str) -> str:
    clean = re.sub(r"\s+", "", _chunk_label(text, ""))
    return re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9]+", "", clean)


def _chapter_titles_match(left: str, right: str) -> bool:
    a = _normalize_chapter_key(left)
    b = _normalize_chapter_key(right)
    if not a or not b:
        return False
    return a == b or a in b or b in a


def _parse_answer_sections(answer_text: str) -> list[dict]:
    clean = str(answer_text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not clean:
        return []
    sections: list[dict] = []
    current_title = "参考答案"
    current_lines: list[str] = []

    def flush() -> None:
        content = "\n".join(current_lines).strip()
        if not content:
            return
        entries = _parse_answer_entries(content)
        if entries:
            sections.append({"title": _chunk_label(current_title, "参考答案"), "answers": entries})

    for raw in clean.splitlines():
        line = raw.strip()
        if not line:
            current_lines.append(raw)
            continue
        if re.match(r"^---\s*第\s*\d+\s*页\s*---$", line) or re.match(r"^\d{1,4}$", line):
            continue
        if re.match(r"^(?:参考答案|答案速查|参考答案及解析|答案解析)\s*$", line):
            continue
        if _is_question_bank_chapter_heading(line):
            flush()
            current_title = line
            current_lines = []
            continue
        current_lines.append(raw)
    flush()
    return sections


def _answers_for_chapter(chapter_title: str, answer_sections: list[dict]) -> dict[str, str]:
    if not answer_sections:
        return {}
    matched = next((section for section in answer_sections if _chapter_titles_match(chapter_title, str(section.get("title") or ""))), None)
    if matched:
        return dict(matched.get("answers") or {})
    if len(answer_sections) == 1:
        return dict(answer_sections[0].get("answers") or {})
    return {}


def _split_question_bank_chunks(text: str, max_chunks: int = _LOCAL_STUDY_QUESTION_MAX_CHUNKS) -> list[dict]:
    body, answers = _split_answer_section(text)
    answer_sections = _parse_answer_sections(answers)
    chapters = _split_question_bank_chapters(body)
    chunks: list[dict] = []

    for chapter in chapters:
        chapter_title = str(chapter.get("title") or "未分章")
        intro, blocks = _split_question_blocks(str(chapter.get("content") or ""))
        answer_lookup = _answers_for_chapter(chapter_title, answer_sections)
        if not blocks:
            continue
        for start in range(0, len(blocks), _LOCAL_STUDY_QUESTION_GROUP_SIZE):
            group = blocks[start : start + _LOCAL_STUDY_QUESTION_GROUP_SIZE]
            if not group:
                continue
            group_numbers = [str(block.get("number") or "").strip() for block in group]
            group_text = "\n\n".join(str(block.get("content") or "").strip() for block in group if str(block.get("content") or "").strip())
            if start == 0 and intro:
                group_text = f"{intro}\n\n{group_text}".strip()
            answer_lines = [f"{number}. {answer_lookup[number]}" for number in group_numbers if number and number in answer_lookup]
            if answer_lines:
                group_text = f"{group_text.rstrip()}\n\n参考答案\n" + "\n".join(answer_lines)
            first = group_numbers[0] if group_numbers else str(start + 1)
            last = group_numbers[-1] if group_numbers else str(start + len(group))
            chunks.append({"label": f"{chapter_title} 第{first}-{last}题", "content": group_text})

    if not chunks:
        return _split_import_chunks(body, max_chunks=max_chunks)
    if max_chunks > 0 and len(chunks) > max_chunks:
        chunks = chunks[:max_chunks]
        chunks[-1]["content"] = (
            str(chunks[-1].get("content") or "").rstrip()
            + "\n\n...[后续题组超过自动拆分上限，先保留前面题组]"
        )
    return chunks


def _split_local_material_chunks(text: str, max_chunks: int = _LOCAL_STUDY_MATERIAL_MAX_CHUNKS) -> list[dict]:
    return _split_import_chunks(text, max_chunks=max_chunks)


def _infer_local_source_type(path: Path, detected_source_type: str) -> str:
    text = "/".join(path.parts[-3:])
    if "题库" in text or "刷题" in path.name or "题" in path.stem:
        return "question_bank"
    return detected_source_type


def _infer_local_module_id(path: Path) -> str:
    text = "/".join(path.parts[-3:])
    rules = (
        ("law", ("法律", "宪法", "民法", "刑法", "行政法", "法理")),
        ("philosophy", ("哲学", "唯物", "辩证法", "认识论", "历史观")),
        ("economy", ("经济", "财政", "货币", "市场", "宏观", "微观")),
        ("writing", ("公文", "写作", "通知", "请示", "报告")),
        ("rural", ("三农", "乡村", "农村", "农业", "农民")),
        ("governance", ("基层", "治理", "群众", "调解", "信访")),
        ("village_affairs", ("村务", "村委会", "村干部", "村民")),
        ("computer", ("计算机", "office", "word", "excel", "wps")),
        ("party", ("党建", "党员", "党章", "党纪")),
        ("current_affairs", ("时政", "两会", "政府工作报告")),
    )
    for module_id, keywords in rules:
        if any(keyword in text for keyword in keywords):
            return module_id
    return "inbox"


def _infer_local_chunk_module_id(path: Path, chunk_label: str, content: str, source_type: str) -> str:
    guessed = r2_store.guess_studyroom_module_id(chunk_label, content, "", source_type)
    if guessed != "inbox":
        return guessed
    return _infer_local_module_id(path)


def _local_study_key(root: Path, path: Path) -> str:
    return path.relative_to(root).as_posix()


def _existing_local_study_keys() -> set[str]:
    keys: set[str] = set()
    for item in r2_store.get_studyroom_data().get("items") or []:
        note = str((item or {}).get("note") or "")
        match = re.search(r"本地 study 导入：([^\n]+)", note)
        if match:
            keys.add(match.group(1).strip())
    return keys


def _iter_local_study_files(root: Path) -> list[Path]:
    return sorted(
        path for path in root.rglob("*")
        if path.is_file()
        and path.name != ".DS_Store"
        and path.suffix.lower() in _LOCAL_STUDY_EXTS
    )


def _with_auto_module(data: dict) -> dict:
    payload = dict(data or {})
    module_id = str(payload.get("module_id") or "").strip() or "inbox"
    if module_id == "inbox":
        payload["module_id"] = r2_store.guess_studyroom_module_id(
            payload.get("title"),
            payload.get("content"),
            payload.get("url"),
            payload.get("source_type"),
        )
    return payload


def register_routes(bp) -> None:
    @bp.route("/studyroom", methods=["GET"])
    def miniapp_studyroom_get():
        return jsonify({"ok": True, "data": r2_store.get_studyroom_data()})

    @bp.route("/studyroom/import", methods=["POST"])
    def miniapp_studyroom_import():
        if request.content_length and request.content_length > _MAX_IMPORT_BYTES:
            return jsonify({"ok": False, "error": "文件太大，先控制在 16MB 内"}), 413
        uploaded = request.files.get("file")
        if not uploaded:
            return jsonify({"ok": False, "error": "没有收到文件"}), 400
        filename = str(uploaded.filename or "资料").strip() or "资料"
        content = uploaded.read() or b""
        if not content:
            return jsonify({"ok": False, "error": "文件是空的"}), 400
        if len(content) > _MAX_IMPORT_BYTES:
            return jsonify({"ok": False, "error": "文件太大，先控制在 16MB 内"}), 413

        try:
            text, detected_source_type = _extract_import_text(filename, content)
        except ValueError as exc:
            return jsonify({"ok": False, "error": str(exc)}), 400
        except RuntimeError as exc:
            return jsonify({"ok": False, "error": str(exc)}), 503
        except Exception as exc:
            return jsonify({"ok": False, "error": f"解析失败：{exc}"}), 400

        if not text:
            return jsonify({"ok": False, "error": "没有抽到文字；如果是扫描版 PDF/图片，需要 OCR，当前还没接。"}), 400

        title = str(request.form.get("title") or "").strip() or Path(filename).stem or "未命名资料"
        module_id = str(request.form.get("module_id") or "inbox").strip() or "inbox"
        requested_source_type = str(request.form.get("source_type") or "").strip()
        source_type = requested_source_type if requested_source_type in _SOURCE_TYPE_OVERRIDES else detected_source_type
        chunks = _split_question_bank_chunks(text) if source_type == "question_bank" else _split_import_chunks(text)
        created_items = []
        for index, chunk in enumerate(chunks, start=1):
            chunk_label = str(chunk.get("label") or "").strip()
            note = f"导入文件：{filename}"
            if len(chunks) > 1:
                note += f"\n自动拆分：第 {index}/{len(chunks)} 段 · {_chunk_label(chunk_label, f'第 {index} 段')}"
            item = r2_store.add_studyroom_item(
                _with_auto_module(
                    {
                        "title": _chunked_title(title, index, len(chunks), chunk_label),
                        "content": str(chunk.get("content") or "").strip(),
                        "module_id": module_id,
                        "source_type": source_type,
                        "status": "todo",
                        "note": note,
                    }
                )
            )
            if item:
                created_items.append(item)
        if not created_items:
            return jsonify({"ok": False, "error": "导入后保存失败"}), 500
        return jsonify(
            {
                "ok": True,
                "item": created_items[0],
                "items": created_items,
                "data": r2_store.get_studyroom_data(),
                "chars": len(text),
                "chunks": len(created_items),
            }
        )

    @bp.route("/studyroom/local-study/import", methods=["POST"])
    def miniapp_studyroom_local_study_import():
        root = _LOCAL_STUDY_ROOT.resolve()
        if not root.exists() or not root.is_dir():
            return jsonify({"ok": False, "error": f"没有找到本地 study 文件夹：{root}", "root": str(root)}), 404

        body = request.get_json(silent=True) or {}
        try:
            max_files = max(1, min(int(body.get("max_files") or 50), 100))
        except Exception:
            max_files = 50
        skip_existing = bool(body.get("skip_existing", True))
        files = _iter_local_study_files(root)[:max_files]
        if not files:
            return jsonify({"ok": False, "error": f"study 文件夹里没有可导入文件：{root}", "root": str(root)}), 404

        existing_keys = _existing_local_study_keys() if skip_existing else set()
        created_items = []
        imported_files = 0
        skipped_files = 0
        errors: list[dict] = []

        for path in files:
            key = _local_study_key(root, path)
            if key in existing_keys:
                skipped_files += 1
                continue
            try:
                text, detected_source_type = _extract_local_file_text(path)
                if not text.strip():
                    skipped_files += 1
                    continue
                source_type = _infer_local_source_type(path, detected_source_type)
                chunks = _split_question_bank_chunks(text) if source_type == "question_bank" else _split_local_material_chunks(text)
                chunks = [chunk for chunk in chunks if str(chunk.get("content") or "").strip()]
                if source_type == "question_bank" and not any("\n\n参考答案\n" in str(chunk.get("content") or "") for chunk in chunks):
                    skipped_files += 1
                    errors.append({"file": key, "error": "题库未识别到可判参考答案，已跳过"})
                    continue
                if not chunks:
                    skipped_files += 1
                    continue
                imported_files += 1
                title = path.stem.strip() or "未命名资料"
                for index, chunk in enumerate(chunks, start=1):
                    chunk_label = str(chunk.get("label") or "").strip()
                    content = str(chunk.get("content") or "").strip()
                    module_id = _infer_local_chunk_module_id(path, chunk_label, content, source_type)
                    item = r2_store.add_studyroom_item(
                        _with_auto_module(
                            {
                                "title": _chunked_title(title, index, len(chunks), chunk_label),
                                "content": content,
                                "module_id": module_id,
                                "source_type": source_type,
                                "status": "todo",
                                "note": "\n".join(
                                    [
                                        f"本地 study 导入：{key}",
                                        f"文件类型：{path.parent.name}",
                                        f"自动拆分：第 {index}/{len(chunks)} 段 · {_chunk_label(chunk_label, f'第 {index} 段')}",
                                    ]
                                ),
                            }
                        )
                    )
                    if item:
                        created_items.append(item)
                existing_keys.add(key)
            except Exception as exc:
                errors.append({"file": key, "error": str(exc)})

        return jsonify(
            {
                "ok": not errors or bool(created_items),
                "root": str(root),
                "files": len(files),
                "imported_files": imported_files,
                "skipped_files": skipped_files,
                "created": len(created_items),
                "errors": errors[:10],
                "data": r2_store.get_studyroom_data(),
            }
        )

    @bp.route("/studyroom/items", methods=["POST"])
    def miniapp_studyroom_add_item():
        data = request.get_json(silent=True) or {}
        item = r2_store.add_studyroom_item(_with_auto_module(data))
        if not item:
            return jsonify({"ok": False, "error": "资料内容不能为空"}), 400
        return jsonify({"ok": True, "item": item, "data": r2_store.get_studyroom_data()})

    @bp.route("/studyroom/items/<item_id>", methods=["PUT"])
    def miniapp_studyroom_update_item(item_id: str):
        data = request.get_json(silent=True) or {}
        item = r2_store.update_studyroom_item(item_id, data)
        if not item:
            return jsonify({"ok": False, "error": "未找到资料或内容无效"}), 404
        return jsonify({"ok": True, "item": item, "data": r2_store.get_studyroom_data()})

    @bp.route("/studyroom/items/<item_id>", methods=["DELETE"])
    def miniapp_studyroom_delete_item(item_id: str):
        ok = r2_store.delete_studyroom_item(item_id)
        if not ok:
            return jsonify({"ok": False, "error": "未找到资料"}), 404
        return jsonify({"ok": True, "data": r2_store.get_studyroom_data()})

    @bp.route("/studyroom/items/<item_id>/codex-sort", methods=["POST"])
    def miniapp_studyroom_codex_sort(item_id: str):
        data = r2_store.get_studyroom_data()
        items = data.get("items") or []
        item = next((x for x in items if str((x or {}).get("id") or "") == str(item_id or "")), None)
        if not item:
            return jsonify({"ok": False, "error": "未找到资料"}), 404
        modules = {str((m or {}).get("id") or ""): str((m or {}).get("label") or "") for m in data.get("modules") or []}
        content_parts = [
            str(item.get("content") or "").strip(),
            str(item.get("note") or "").strip(),
            str(item.get("url") or "").strip(),
        ]
        content = "\n\n".join([x for x in content_parts if x]).strip()
        if not content:
            return jsonify({"ok": False, "error": "这条资料没有可整理内容"}), 400
        task = codex_group_chat.create_task(
            {
                "mode": "studyroom",
                "window_id": "studyroom",
                "reply_target": "studyroom",
                "study_item_id": item_id,
                "study_title": item.get("title") or "",
                "study_module": modules.get(str(item.get("module_id") or ""), "待整理"),
                "study_source": item.get("source_type") or "",
                "study_url": item.get("url") or "",
                "user_message": content,
                "client_request_id": f"studyroom-{item_id}",
            }
        )
        if not task:
            return jsonify({"ok": False, "error": "创建整理任务失败"}), 500
        r2_store.update_studyroom_item(item_id, {"status": "sorting"})
        return jsonify({"ok": True, "task": task, "data": r2_store.get_studyroom_data()})

    @bp.route("/studyroom/items/<item_id>/auto-module", methods=["POST"])
    def miniapp_studyroom_auto_module(item_id: str):
        data = r2_store.get_studyroom_data()
        items = data.get("items") or []
        item = next((x for x in items if str((x or {}).get("id") or "") == str(item_id or "")), None)
        if not item:
            return jsonify({"ok": False, "error": "未找到资料"}), 404
        module_id = r2_store.guess_studyroom_module_id(
            item.get("title"),
            item.get("content"),
            item.get("url"),
            item.get("source_type"),
        )
        updated = r2_store.update_studyroom_item(item_id, {"module_id": module_id})
        if not updated:
            return jsonify({"ok": False, "error": "自动归类失败"}), 500
        return jsonify({"ok": True, "item": updated, "module_id": module_id, "data": r2_store.get_studyroom_data()})

    @bp.route("/studyroom/study-logs", methods=["POST"])
    def miniapp_studyroom_add_study_log():
        data = request.get_json(silent=True) or {}
        entry = r2_store.add_studyroom_log(str(data.get("content") or ""))
        if not entry:
            return jsonify({"ok": False, "error": "学习记录不能为空"}), 400
        return jsonify({"ok": True, "entry": entry, "data": r2_store.get_studyroom_data()})
